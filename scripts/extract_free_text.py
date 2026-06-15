#!/usr/bin/env python3
"""
Extract plain text from comment attachments (PDF/DOCX) with structure preserved
for LLM-friendly analysis. Writes fullText_1, fullText_2, ... into the comment
records and updates data/all_comments.json and data/comments.csv.

Usage:
  python extract_free_text.py [--limit N] [--dry-run]
  python extract_free_text.py --clear-fulltext   # remove all fullText_* and write clean JSON/CSV
  python extract_free_text.py [--resume]         # skip comments that already have fullText (default: resume)

Reads data/all_comments.json and data/attachments/<documentId>/attachment_*.pdf|.docx.
Saves JSON/CSV after each comment so interruption can be resumed.
"""

from __future__ import annotations

import argparse
import json
import logging
import sys
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument

# Suppress only "CropBox missing from /Page, defaulting to MediaBox" (often printed to stderr)
class _CropBoxStderrFilter:
    """File-like that forwards stderr but drops lines containing CropBox/defaulting to MediaBox."""
    def __init__(self, stream):
        self._stream = stream
        self._buf = ""

    def write(self, s: str) -> int:
        if not isinstance(s, str):
            return 0
        self._buf += s
        while "\n" in self._buf or "\r" in self._buf:
            sep = "\n" if "\n" in self._buf else "\r"
            line, _, self._buf = self._buf.partition(sep)
            if "CropBox" in line and "defaulting to MediaBox" in line:
                continue
            self._stream.write(line + sep)
        return len(s)

    def flush(self) -> None:
        if self._buf:
            if "CropBox" not in self._buf or "defaulting to MediaBox" not in self._buf:
                self._stream.write(self._buf)
            self._buf = ""
        self._stream.flush()


class _CropBoxLogFilter(logging.Filter):
    def filter(self, record: logging.LogRecord) -> bool:
        msg = (record.getMessage() or "")
        if "CropBox" in msg and "defaulting to MediaBox" in msg:
            return False
        return True


def _suppress_cropbox_warning() -> None:
    for name in ("pypdf", "pdfminer"):
        logging.getLogger(name).addFilter(_CropBoxLogFilter())
    # Root logger so any child (e.g. pdfminer.pdfpage) is also filtered for this message
    root = logging.getLogger()
    root.addFilter(_CropBoxLogFilter())
    # Stderr: library may print() this message
    if hasattr(sys, "stderr") and not isinstance(sys.stderr, _CropBoxStderrFilter):
        sys.stderr = _CropBoxStderrFilter(sys.stderr)


_suppress_cropbox_warning()

# Reuse data paths from pull_comments
DATA_DIR = Path("data")
COMMENTS_JSON_FILE = DATA_DIR / "all_comments.json"
COMMENTS_CSV_FILE = DATA_DIR / "comments.csv"
ATTACHMENTS_DIR = DATA_DIR / "attachments"
EXTRACT_LOG_FILE = DATA_DIR / "extract_free_text.log"

# Extensions to try per attachment (same order as download: primary is pdf or docx)
ATTACHMENT_EXTS = ("pdf", "docx")


def _setup_logging() -> logging.Logger:
    log = logging.getLogger("extract_free_text")
    log.setLevel(logging.DEBUG)
    if log.handlers:
        return log
    fmt = logging.Formatter("%(asctime)s %(levelname)s %(message)s")
    out = logging.StreamHandler(sys.stdout)
    out.setFormatter(fmt)
    log.addHandler(out)
    if DATA_DIR.exists():
        try:
            fh = logging.FileHandler(EXTRACT_LOG_FILE, encoding="utf-8")
            fh.setFormatter(fmt)
            log.addHandler(fh)
        except OSError:
            pass
    return log


logger = _setup_logging()


def _table_to_markdown(table: list[list[str]]) -> str:
    """Turn a grid of cells into a markdown table."""
    if not table:
        return ""
    rows = [[str(c or "").strip().replace("\n", " ") for c in row] for row in table]
    lines = []
    for i, row in enumerate(rows):
        lines.append("| " + " | ".join(row) + " |")
        if i == 0:
            lines.append("| " + " | ".join("---" for _ in row) + " |")
    return "\n".join(lines)


def extract_pdf(path: Path) -> str:
    """Extract text from a PDF with layout and tables, formatted for LLM consumption."""
    parts = []
    try:
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                page_parts = []
                # Layout-preserving text (paragraphs, line breaks)
                text = page.extract_text(layout=True)
                if text and text.strip():
                    page_parts.append(text.strip())
                # Tables as markdown
                tables = page.find_tables()
                for t in tables:
                    extracted = t.extract()
                    if extracted:
                        md = _table_to_markdown(extracted)
                        if md:
                            page_parts.append("\n\n" + md)
                page_text = "\n\n".join(p for p in page_parts if p)
                if page_text:
                    parts.append(f"--- Page {i + 1} ---\n\n{page_text}")
    except Exception as e:
        logger.warning("PDF extraction failed for %s: %s", path, e)
        return f"[PDF extraction error: {e}]"
    return "\n\n".join(parts) if parts else ""


def extract_docx(path: Path) -> str:
    """Extract text from a DOCX with headings and tables, formatted for LLM consumption."""
    parts = []
    try:
        doc = DocxDocument(path)
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = (para.style and para.style.name or "").lower()
            if "heading 1" in style:
                parts.append("\n\n## " + text + "\n")
            elif "heading 2" in style:
                parts.append("\n\n### " + text + "\n")
            elif "heading" in style:
                parts.append("\n\n#### " + text + "\n")
            else:
                parts.append(text)
        for table in doc.tables:
            rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
            if rows:
                md = _table_to_markdown(rows)
                if md:
                    parts.append("\n\n" + md + "\n")
    except Exception as e:
        logger.warning("DOCX extraction failed for %s: %s", path, e)
        return f"[DOCX extraction error: {e}]"
    return "\n\n".join(p.strip() for p in parts if p.strip()).strip()


def _find_attachment_path(doc_id: str, att_count: int) -> tuple[Path | None, str | None]:
    """Return (path, ext) for attachment_<att_count> or (None, None) if not found."""
    out_dir = ATTACHMENTS_DIR / doc_id
    for ext in ATTACHMENT_EXTS:
        p = out_dir / f"attachment_{att_count}.{ext}"
        if p.exists() and p.stat().st_size > 0:
            return (p, ext)
    return (None, None)


def extract_attachment(path: Path, ext: str) -> str:
    if ext == "pdf":
        return extract_pdf(path)
    if ext == "docx":
        return extract_docx(path)
    return f"[Unsupported format: .{ext}]"


def _is_extracted(text: str) -> bool:
    """True if fullText value counts as already extracted (resume skip)."""
    if not text or not isinstance(text, str):
        return False
    # Error/placeholder values start with "["
    return not text.strip().startswith("[")


def _save_progress(comments: list[dict]) -> None:
    """Write current comments to JSON and CSV (checkpoint)."""
    with open(COMMENTS_JSON_FILE, "w", encoding="utf-8") as f:
        json.dump(comments, f, indent=2, ensure_ascii=False)
    from pull_comments import write_csv
    write_csv(comments, COMMENTS_CSV_FILE)


def run_extraction(
    comments: list[dict],
    limit: int | None = None,
    dry_run: bool = False,
    resume: bool = True,
    checkpoint_every_comment: bool = True,
) -> None:
    """Add fullText_1, fullText_2, ... to each comment from attachment files.
    Writes JSON/CSV after each comment so interrupt can be resumed (resume=True skips already-extracted).
    Never writes partial PDF text: we only set fullText after extract_attachment() returns.
    """
    if limit is not None:
        comments = comments[:limit]
    logger.info("Processing %d comments (resume=%s)", len(comments), resume)
    total_attachments = 0
    extracted = 0
    skipped = 0
    missing = 0
    errors = 0
    comment_index = 0
    for item in comments:
        data = item.get("data", {})
        doc_id = data.get("id") or (data.get("attributes") or {}).get("documentId")
        if not doc_id:
            comment_index += 1
            continue
        att_count = 0
        comment_dirty = False
        for idx, inc in enumerate(item.get("included", [])):
            if inc.get("type") != "attachments":
                continue
            att_count += 1
            total_attachments += 1
            n = idx + 1  # same index as CSV (included position)
            path, ext = _find_attachment_path(doc_id, att_count)
            if path is None:
                if not resume or not _is_extracted(item.get(f"fullText_{n}", "")):
                    item[f"fullText_{n}"] = ""
                    comment_dirty = True
                missing += 1
                if total_attachments <= 3 or missing <= 5:
                    logger.debug("No file for %s attachment_%d", doc_id, att_count)
                continue
            if dry_run:
                item[f"fullText_{n}"] = f"[dry-run: would extract {path.name}]"
                extracted += 1
                comment_dirty = True
                continue
            if resume and _is_extracted(item.get(f"fullText_{n}", "")):
                skipped += 1
                continue
            try:
                logger.info("Extracting %s attachment_%d (%s)", doc_id, att_count, path.name)
                text = extract_attachment(path, ext)
                item[f"fullText_{n}"] = text
                extracted += 1
                comment_dirty = True
            except Exception as e:
                logger.warning("Extract failed %s attachment_%d: %s", doc_id, att_count, e)
                item[f"fullText_{n}"] = f"[Extraction error: {e}]"
                errors += 1
                comment_dirty = True
        if checkpoint_every_comment and comment_dirty and not dry_run:
            _save_progress(comments)
            logger.debug("Checkpoint: saved after comment %s", doc_id)
        comment_index += 1
    logger.info(
        "Attachments: %d total, %d extracted, %d skipped (resume), %d missing file, %d errors",
        total_attachments,
        extracted,
        skipped,
        missing,
        errors,
    )


def clear_fulltext(comments: list[dict]) -> None:
    """Remove all fullText_* keys from every comment; write JSON and CSV."""
    removed = 0
    for item in comments:
        keys = [k for k in item if isinstance(k, str) and k.startswith("fullText_")]
        for k in keys:
            del item[k]
            removed += 1
    logger.info("Removed %d fullText_* keys from comments", removed)
    _save_progress(comments)
    logger.info("Wrote %s and %s (clean, no fullText)", COMMENTS_JSON_FILE, COMMENTS_CSV_FILE)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__.strip().split("\n")[1])
    parser.add_argument("--limit", type=int, default=None, help="Process only first N comments (for testing)")
    parser.add_argument("--dry-run", action="store_true", help="Do not write files; log what would be done")
    parser.add_argument("--clear-fulltext", action="store_true", help="Remove all fullText_* from JSON/CSV and exit (start clean)")
    parser.add_argument("--no-resume", action="store_true", help="Re-extract all (ignore existing fullText)")
    args = parser.parse_args()

    if not COMMENTS_JSON_FILE.exists():
        logger.error("%s not found. Run pull_comments.py first.", COMMENTS_JSON_FILE)
        sys.exit(1)

    with open(COMMENTS_JSON_FILE, encoding="utf-8") as f:
        comments = json.load(f)

    if args.clear_fulltext:
        clear_fulltext(comments)
        return

    run_extraction(
        comments,
        limit=args.limit,
        dry_run=args.dry_run,
        resume=not args.no_resume,
        checkpoint_every_comment=True,
    )

    if args.dry_run:
        logger.info("Dry run: not writing JSON or CSV")
        return

    _save_progress(comments)
    logger.info("Wrote %s and %s", COMMENTS_JSON_FILE, COMMENTS_CSV_FILE)


if __name__ == "__main__":
    main()
