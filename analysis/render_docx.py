"""Build manuscript.docx from manuscript.md.

Renders Markdown headings, paragraphs, lists, tables, and inline emphasis into
a JAMA-style Word document. Inserts the manuscript figures as embedded images
where the manuscript text references them. Footnote-style superscript citations
(<sup>N</sup>) are rendered as superscript runs.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "manuscript.md"
DOCX_PATH = ROOT / "manuscript.docx"

# Map figure captions -> on-disk image paths (v0.5: stakeholder × axis lead)
FIGURE_PATHS = {
    "Figure 1": ROOT / "output" / "coalitions" / "fig_coalition_length_density.png",
    "Figure 2": ROOT / "output" / "stakeholder_positions" / "fig_stakeholder_modal_heatmap.png",
    "Figure 3": ROOT / "output" / "stakeholder_positions" / "fig_industry_vs_patient.png",
    "Figure 4": ROOT / "output" / "coalitions" / "fig_coalition_pca.png",
    "Figure 5": ROOT / "output" / "rfi_coverage" / "fig_rfi_x_coalition.png",
    "Figure S1": ROOT / "output" / "manuscript" / "figure_prisma_flow.png",
    "Figure S2": ROOT / "output" / "manuscript" / "figure_irr_forest.png",
    "Figure S3": ROOT / "output" / "cluster_validation" / "fig_cluster_validation.png",
    "Figure S4": ROOT / "output" / "stakeholder_positions" / "fig_stakeholder_position_profile.png",
    "Figure S5": ROOT / "output" / "coalitions" / "fig_coalition_topic_emphasis.png",
}


def set_doc_styles(doc: Document) -> None:
    """JAMA-ish defaults: Times New Roman 11, 1.5 line spacing on body."""
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    for h_style, size, bold in [("Heading 1", 16, True), ("Heading 2", 13, True),
                                 ("Heading 3", 12, True)]:
        s = styles[h_style]
        s.font.name = "Times New Roman"
        s.font.size = Pt(size)
        s.font.bold = bold
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after = Pt(6)


def add_run_segments(p, text: str) -> None:
    """Render inline Markdown (*italic*, **bold**, `code`, <sup>N</sup>) into a paragraph."""
    # Tokenize on the four inline patterns
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|<sup>[^<]+</sup>|<sub>[^<]+</sub>)"
    )
    for chunk in pattern.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            r = p.add_run(chunk[2:-2])
            r.bold = True
        elif chunk.startswith("*") and chunk.endswith("*") and len(chunk) > 2:
            r = p.add_run(chunk[1:-1])
            r.italic = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            r = p.add_run(chunk[1:-1])
            r.font.name = "Courier New"
        elif chunk.startswith("<sup>") and chunk.endswith("</sup>"):
            r = p.add_run(chunk[5:-6])
            r.font.superscript = True
        elif chunk.startswith("<sub>") and chunk.endswith("</sub>"):
            r = p.add_run(chunk[5:-6])
            r.font.subscript = True
        else:
            p.add_run(chunk)


def add_image(doc: Document, fig_label: str, caption_text: str) -> None:
    """Insert image then caption."""
    path = FIGURE_PATHS.get(fig_label)
    if path and path.exists():
        doc.add_picture(str(path), width=Inches(6.5))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_p.add_run(f"{fig_label}. ")
    cap_run.bold = True
    add_run_segments(cap_p, caption_text)


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    """Read a contiguous Markdown table starting at lines[start]. Returns (rows, end_idx)."""
    rows = []
    i = start
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        line = lines[i].strip()
        # Skip separator row
        if re.match(r"^\|[\s:|-]+\|$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def write_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(rows):
        for c_idx, cell in enumerate(row):
            tc = table.rows[r_idx].cells[c_idx]
            tc.text = ""
            p = tc.paragraphs[0]
            add_run_segments(p, cell)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True


def render_markdown(md_text: str, doc: Document) -> None:
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == "---":
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue

        # Tables
        if line.lstrip().startswith("|"):
            rows, i = parse_markdown_table(lines, i)
            write_table(doc, rows)
            continue

        # Block-quote
        if line.startswith(">"):
            quote = line.lstrip("> ").strip()
            j = i + 1
            while j < len(lines) and lines[j].startswith(">"):
                quote += " " + lines[j].lstrip("> ").strip()
                j += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            r = p.add_run(quote)
            r.italic = True
            i = j
            continue

        # Numbered list
        m = re.match(r"^\s*(\d+)\.\s+(.*)$", line)
        if m:
            text = m.group(2)
            p = doc.add_paragraph(style="List Number")
            add_run_segments(p, text)
            i += 1
            continue

        # Bulleted list
        if re.match(r"^\s*[-*]\s+", line):
            text = re.sub(r"^\s*[-*]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            add_run_segments(p, text)
            i += 1
            continue

        # Figure caption shorthand: lines like **Figure N.** ... (with or without [→ path])
        fig_match = re.match(r"^\*\*Figure (\w+)\.\*\*\s+(.+?)\s*$", line)
        if fig_match:
            label = f"Figure {fig_match.group(1)}"
            caption = re.sub(r"\s*\[→[^\]]+\]\s*", "", fig_match.group(2)).strip()
            add_image(doc, label, caption)
            i += 1
            continue

        # Strip table-locator phrases like "[→ output/...]" leaving just the prose
        line = re.sub(r"\s*\[→[^\]]+\]\s*", "", line)

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Plain paragraph (could span multiple lines until blank)
        para_lines = [line]
        j = i + 1
        while j < len(lines) and lines[j].strip() and not (
            lines[j].startswith("#") or lines[j].lstrip().startswith("|") or
            lines[j].startswith(">") or
            re.match(r"^\s*\d+\.\s+", lines[j]) or
            re.match(r"^\s*[-*]\s+", lines[j]) or
            lines[j].startswith("```") or lines[j].strip() == "---"
        ):
            para_lines.append(lines[j])
            j += 1
        para_text = " ".join(para_lines).strip()
        para_text = re.sub(r"\s*\[→[^\]]+\]\s*", "", para_text)
        if para_text:
            p = doc.add_paragraph()
            add_run_segments(p, para_text)
        i = j


def main() -> None:
    md = MD_PATH.read_text(encoding="utf-8")
    doc = Document()
    set_doc_styles(doc)
    render_markdown(md, doc)
    doc.save(str(DOCX_PATH))
    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    main()
